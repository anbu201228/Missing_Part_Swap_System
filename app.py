from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import os
from datetime import datetime
from openpyxl import Workbook
from docx import Document
import io

app = Flask(__name__)
app.secret_key = "alstom_secret_key" # Required for flashing messages

# Configuration
UPLOAD_FOLDER = 'static/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Ensure upload folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

db = SQLAlchemy(app)

# Database Model
class SwapRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    requester_name = db.Column(db.String(100), nullable=False)
    project = db.Column(db.String(50))
    reason = db.Column(db.Text)
    installation_name = db.Column(db.String(100))
    part_no = db.Column(db.String(50))
    part_name = db.Column(db.String(100))
    from_ts = db.Column(db.String(50))
    fix_ts = db.Column(db.String(50))
    image = db.Column(db.String(100))
    date = db.Column(db.String(20))
    status = db.Column(db.String(20), default="Pending")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/new', methods=['GET', 'POST'])
def new_request():
    if request.method == 'POST':
        try:
            # Secure Image Upload Handling
            image_name = ""
            if 'image' in request.files:
                file = request.files['image']
                if file and file.filename != '':
                    filename = secure_filename(file.filename)
                    # Prepend timestamp to filename to prevent overwrites
                    image_name = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], image_name))

            # Database Entry
            new_entry = SwapRequest(
                requester_name=request.form.get('requester_name'),
                project=request.form.get('project'),
                reason=request.form.get('reason'),
                installation_name=request.form.get('installation_name'),
                part_no=request.form.get('part_no'),
                part_name=request.form.get('part_name'),
                from_ts=request.form.get('from_ts'),
                fix_ts=request.form.get('fix_ts'),
                image=image_name,
                date=request.form.get('date') if request.form.get('date') else datetime.now().strftime('%Y-%m-%d')
            )

            db.session.add(new_entry)
            db.session.commit()
            flash("Request submitted successfully!", "success")
            return redirect(url_for('view_requests'))
            
        except Exception as e:
            db.session.rollback()
            flash(f"Error: {str(e)}", "danger")
            return redirect(url_for('new_request'))

    return render_template('new_request.html')

@app.route('/requests')
def view_requests():
    # Show newest requests first
    data = SwapRequest.query.all()
    return render_template('view_requests.html', data=data)

@app.route('/request/<int:id>')
def view_request(id):
    req = SwapRequest.query.get_or_404(id)
    return render_template('request_detail.html', request=req)

@app.route('/edit/<int:id>', methods=['GET', 'POST'])
def edit_request(id):
    req = SwapRequest.query.get_or_404(id)
    if request.method == 'POST':
        try:
            # Handle image upload if new image
            image_name = req.image
            if 'image' in request.files:
                file = request.files['image']
                if file and file.filename != '':
                    filename = secure_filename(file.filename)
                    image_name = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], image_name))
                    # Optionally delete old image
                    if req.image and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], req.image)):
                        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], req.image))

            req.requester_name = request.form.get('requester_name')
            req.project = request.form.get('project')
            req.reason = request.form.get('reason')
            req.installation_name = request.form.get('installation_name')
            req.part_no = request.form.get('part_no')
            req.part_name = request.form.get('part_name')
            req.from_ts = request.form.get('from_ts')
            req.fix_ts = request.form.get('fix_ts')
            req.image = image_name
            req.date = request.form.get('date') if request.form.get('date') else req.date
            req.status = request.form.get('status')

            db.session.commit()
            flash("Request updated successfully!", "success")
            return redirect(url_for('view_requests'))
        except Exception as e:
            db.session.rollback()
            flash(f"Error: {str(e)}", "danger")
            return redirect(url_for('edit_request', id=id))

    return render_template('edit_request.html', request=req)

@app.route('/delete/<int:id>')
def delete_request(id):
    req = SwapRequest.query.get_or_404(id)
    try:
        # Delete image if exists
        if req.image and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], req.image)):
            os.remove(os.path.join(app.config['UPLOAD_FOLDER'], req.image))
        db.session.delete(req)
        db.session.commit()
        flash("Request deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error: {str(e)}", "danger")
    return redirect(url_for('view_requests'))

@app.route('/export/excel')
def export_excel():
    from openpyxl.drawing.image import Image as XLImage
    data = SwapRequest.query.all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Swap Requests"
    
    # Headers
    headers = ['ID', 'Date', 'Requester Name', 'Project', 'Reason', 'Installation Name', 'Part No', 'Part Name', 'From TS', 'Fix TS', 'Status', 'Evidence Image']
    for col_num, header in enumerate(headers, 1):
        ws.cell(row=1, column=col_num, value=header)
    
    # Set column width for image column
    ws.column_dimensions['L'].width = 20
    
    # Data
    for row_num, req in enumerate(data, 2):
        ws.cell(row=row_num, column=1, value=req.id)
        ws.cell(row=row_num, column=2, value=req.date)
        ws.cell(row=row_num, column=3, value=req.requester_name)
        ws.cell(row=row_num, column=4, value=req.project)
        ws.cell(row=row_num, column=5, value=req.reason)
        ws.cell(row=row_num, column=6, value=req.installation_name)
        ws.cell(row=row_num, column=7, value=req.part_no)
        ws.cell(row=row_num, column=8, value=req.part_name)
        ws.cell(row=row_num, column=9, value=req.from_ts)
        ws.cell(row=row_num, column=10, value=req.fix_ts)
        ws.cell(row=row_num, column=11, value=req.status)
        
        # Add image if exists
        if req.image:
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], req.image)
            if os.path.exists(image_path):
                try:
                    img = XLImage(image_path)
                    img.width = 100
                    img.height = 100
                    ws.add_image(img, f'L{row_num}')
                    ws.row_dimensions[row_num].height = 100
                except Exception as e:
                    ws.cell(row=row_num, column=12, value=f'[Image error: {str(e)}]')
            else:
                ws.cell(row=row_num, column=12, value='[Image not found]')
        else:
            ws.cell(row=row_num, column=12, value='[No image]')
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return send_file(output, download_name='swap_requests.xlsx', as_attachment=True)

@app.route('/export/word')
def export_word():
    data = SwapRequest.query.all()
    
    doc = Document()
    doc.add_heading('Swap Requests Report', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    for req in data:
        doc.add_heading(f'Request ID: {req.id}', level=1)
        doc.add_paragraph(f'Date: {req.date}')
        doc.add_paragraph(f'Requester Name: {req.requester_name}')
        doc.add_paragraph(f'Project: {req.project}')
        doc.add_paragraph(f'Reason: {req.reason}')
        doc.add_paragraph(f'Installation Name: {req.installation_name}')
        doc.add_paragraph(f'Part No: {req.part_no}')
        doc.add_paragraph(f'Part Name: {req.part_name}')
        doc.add_paragraph(f'From TS: {req.from_ts}')
        doc.add_paragraph(f'Fix TS: {req.fix_ts}')
        doc.add_paragraph(f'Status: {req.status}')
        
        # Add image if exists
        if req.image:
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], req.image)
            if os.path.exists(image_path):
                try:
                    doc.add_heading('Evidence Image', level=2)
                    doc.add_picture(image_path, width=3000000)  # Width in EMUs
                except Exception as e:
                    doc.add_paragraph(f'[Image could not be embedded: {str(e)}]')
            else:
                doc.add_paragraph('[Image file not found]')
        
        doc.add_paragraph('')  # Empty line
        doc.add_paragraph('_' * 80)  # Separator
        doc.add_paragraph('')
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, download_name='swap_requests.docx', as_attachment=True)

@app.route('/download/image/<filename>')
def download_image(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            flash("Image not found", "danger")
            return redirect(url_for('view_requests'))
    except Exception as e:
        flash(f"Error downloading image: {str(e)}", "danger")
        return redirect(url_for('view_requests'))

@app.route('/download/request/<int:id>')
def download_request(id):
    try:
        req = SwapRequest.query.get_or_404(id)
        
        doc = Document()
        doc.add_heading(f'Request Details - ID: {req.id}', 0)
        doc.add_paragraph('')
        
        doc.add_heading('Information', level=1)
        doc.add_paragraph(f'Date: {req.date}')
        doc.add_paragraph(f'Requester Name: {req.requester_name}')
        doc.add_paragraph(f'Project: {req.project}')
        doc.add_paragraph(f'Reason: {req.reason}')
        doc.add_paragraph(f'Installation Name: {req.installation_name}')
        doc.add_paragraph(f'Part No: {req.part_no}')
        doc.add_paragraph(f'Part Name: {req.part_name}')
        doc.add_paragraph(f'From TS: {req.from_ts}')
        doc.add_paragraph(f'Fix TS: {req.fix_ts}')
        doc.add_paragraph(f'Status: {req.status}')
        doc.add_paragraph('')
        
        if req.image:
            try:
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], req.image)
                if os.path.exists(image_path):
                    doc.add_heading('Evidence Image', level=1)
                    doc.add_picture(image_path, width=3000000)  # Width in EMUs
            except Exception as e:
                doc.add_paragraph(f'[Image could not be embedded: {str(e)}]')
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(output, download_name=f'request_{req.id}.docx', as_attachment=True)
    except Exception as e:
        flash(f"Error downloading request: {str(e)}", "danger")
        return redirect(url_for('view_requests'))

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)